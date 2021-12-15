from docx import Document
# create an instance of a word document we want to open
doc = Document(rba_path)

def collectFromEndnote():
    findBegin = False
    referenceList = []

    for i in range(0, len(doc.paragraphs)):
        if (doc.paragraphs[i].text.lower().startswith('references') or \
            doc.paragraphs[i].text.lower().startswith('referen')):
            for run in doc.paragraphs[i].runs:
                if run.bold:
                    findBegin = True
                    
        if (findBegin == True):
            ref = doc.paragraphs[i].text.strip()
            if (ref.lower().startswith('referen')):
                continue
            elif (ref.strip() == ''):
                continue
            else: 
                
                search_item = {
                    'original_text': ref,
                    'p_authors': '',
                    'p_title': '',
                    'p_doi': ''
                }
                x1 = re.findall("(.*et al.?)([^.]+).", ref)
                x2 = re.findall("(.*)\d{4}.?\s?(\".*\")", ref)
                x3 = re.findall(r'doi:?\s?(.+).?', ref)
                
                if x3:
                    search_item['p_doi'] = x3[0].strip(punctuation).strip()

                if x1:
                    search_item['p_authors'] = x1[0][0].strip(punctuation).strip()
                    search_item['p_title'] = x1[0][1].strip(punctuation).strip()

                elif x2:
                    search_item['p_authors'] = x2[0][0].strip(punctuation).strip()
                    search_item['p_title'] = x2[0][1].strip(punctuation).strip()

                elif (ref.lower().startswith('who') or ref.lower().startswith('lci') \
                      or ref.lower().startswith('nvn') or ref.lower().startswith('nice')):
                    continue
    
    
            referenceList.append(search_item)
    return referenceList