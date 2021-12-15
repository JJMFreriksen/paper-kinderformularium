from Bio import Entrez
email='simgeekiz48@gmail.com'
from Bio.Entrez import efetch, read
import os
from string import punctuation
import re
import csv 

from docx import Document

def search_API(query):
    Entrez.email = email
    handle = Entrez.esearch(db='pubmed', 
                            sort='relevance', 
                            retmax='20',
                            retmode='xml', 
                            term=query)
    results = Entrez.read(handle)
    return results

def fetch_details(id_list):
    ids = ','.join(id_list)
    Entrez.email = email
    handle = Entrez.efetch(db='pubmed',
                           retmode='xml',
                           id=ids)
    results = Entrez.read(handle)
    return results
    
def create_csv_path(rba_path):
    csv_dir = os.path.abspath(os.path.join(os.path.dirname(rba_path), 'csv'))
    if not os.path.exists(csv_dir):
        os.makedirs(csv_dir)

    head, filename = os.path.split(rba_path)
    csv_path = os.path.join(csv_dir, filename.replace('.docx', '.csv'))
    return csv_path
    
def collectFromEndnote(rba_path):
    doc = Document(rba_path) # create an instance of a word document we want to open
    
    findBegin = False
    referenceList = []

    for i in range(0, len(doc.paragraphs)):
        if (doc.paragraphs[i].text.lower().startswith('references') or \
            doc.paragraphs[i].text.lower().startswith('referen')):
            for run in doc.paragraphs[i].runs:
                if run.bold:
                    findBegin = True
                    
        if (findBegin == True):
            ref = doc.paragraphs[i].text.strip()
            if (ref.lower().startswith('referen')):
                continue
            elif (ref.strip() == ''):
                continue
            else: 
                
                search_item = {
                    'original_text': ref,
                    'p_authors': '',
                    'p_title': '',
                    'p_doi': ''
                }
                x1 = re.findall("(.*et al.?)([^.]+).", ref)
                x2 = re.findall("(.*)\d{4}.?\s?(\".*\")", ref)
                x3 = re.findall(r'doi:?\s?(.+).?', ref)
                
                if x3:
                    search_item['p_doi'] = x3[0].strip(punctuation).strip()

                if x1:
                    search_item['p_authors'] = x1[0][0].strip(punctuation).strip()
                    search_item['p_title'] = x1[0][1].strip(punctuation).strip()

                elif x2:
                    search_item['p_authors'] = x2[0][0].strip(punctuation).strip()
                    search_item['p_title'] = x2[0][1].strip(punctuation).strip()

                elif (ref.lower().startswith('who') or ref.lower().startswith('lci') \
                      or ref.lower().startswith('nvn') or ref.lower().startswith('nice')):
                    continue
    
    
            referenceList.append(search_item)
    return referenceList

def collectFromTables(rba_path):
    doc = Document(rba_path)
    # table_headers = ['bron', 'bewijs', 'effect', 'opmerkingen', 'source', 'evidence', 'effect', 'remarks']

    data = []
    referencesFromTables = []
    for table in doc.tables:
        
        for i, row in enumerate(table.rows):
            text = [cell.text.lower() for cell in row.cells]
            if i == 0:
                effect_index = text.index('effect') if 'effect' in text else None
                continue

            for j, cell in enumerate(row.cells):

                if j == effect_index:

                    bold_text = ''
                    p_title = ''
                    p_authors = ''
                    untilSummary = False

                    for p, paragraph in enumerate(cell.paragraphs):

                        if len(paragraph.runs) < 1:
                            continue

                        if (paragraph.runs[0].bold 
                            and paragraph.runs[-1].bold and p_title == ''):

                            bold_text += paragraph.text.strip() + ' '

                        elif (not paragraph.runs[-1].bold 
                              and p > 0
                              and len(cell.paragraphs[p-1].runs) > 0
                              and cell.paragraphs[p-1].runs[-1].bold):

                            p_title = bold_text

                            p_authors = ' '.join([p_authors, paragraph.text.strip()])

                        elif (paragraph.runs[0].underline 
                              and (paragraph.runs[0].text.lower().startswith('samenvatting') 
                                   or paragraph.runs[0].text.lower().startswith('summary'))):

                            break

                    if p_title.strip() != ''  and re.search(r'^[0-9].[0-9]', p_title) is None: 
                        referencesFromTables.append({
                            'p_title': p_title.strip().replace('\xa0',' '),
                            'p_authors': p_authors,
                        })
    return referencesFromTables
