from Bio import Entrez
from Bio.Entrez import efetch, read
import os
from string import punctuation
import re
import csv
from simhash import Simhash, SimhashIndex
from docx import Document

email = 'simgeekiz48@gmail.com'

def search_API(query):
    """Retrieves ids of relevant papers from PubMed API

    Parameters
    ----------
    query : string
        DOI or title from a targeted paper

    Returns
    -------
    results : list of integers
        Plain list of paper ids
    """
    Entrez.email = email
    handle = Entrez.esearch(db='pubmed',
                            sort='relevance',
                            retmax='20',
                            retmode='xml',
                            term=query)
    results = Entrez.read(handle)
    return results

def fetch_details(id_list):
    """Retrieves the papers from PubMed API by their ids

    Parameters
    ----------
    id_list : list of integers

    Returns
    -------
    results : list of dictionaries
    """
    ids = ','.join(id_list)
    Entrez.email = email
    handle = Entrez.efetch(db='pubmed',
                           retmode='xml',
                           id=ids)
    results = Entrez.read(handle)
    return results

def create_csv_path(rba_path):
    """Creates a csv path for the output

    Parameters
    ----------
    rba_path : string
        Path to the file to be processed.
        This file must be a .docx formatted file.

    Returns
    -------
    csv_path : string
    """
    csv_dir = os.path.abspath(os.path.join(os.path.dirname(rba_path), 'csv'))
    if not os.path.exists(csv_dir):
        os.makedirs(csv_dir)

    head, filename = os.path.split(rba_path)
    csv_path = os.path.join(csv_dir, filename.replace('.docx', '.csv'))
    return csv_path

def collectFromEndnote(rba_path):
    """Extracts the reference list from a file with Endnote

    Parameters
    ----------
    rba_path : string
        Path to the file to be processed.
        This file must be a .docx formatted file.
        The file must contain a text sayin at least "referen"
        in order to function properly

    Returns
    -------
    references_list : list of dictionaries
        Contains the extracted DOI, authors, and title
        for each reference as an object.
    """
    doc = Document(rba_path) # create an instance of a word document we want to open

    findBegin = False
    references_list = []

    for i in range(0, len(doc.paragraphs)):
        if (doc.paragraphs[i].text.lower().startswith('references') or \
            doc.paragraphs[i].text.lower().startswith('referen')):
            for run in doc.paragraphs[i].runs:
                if run.bold:
                    findBegin = True

        if (findBegin == True):
            ref = doc.paragraphs[i].text.strip()
            if (ref.lower().startswith('referen')):
                continue
            elif (ref.strip() == ''):
                continue
            else:

                search_item = {
                    'original_text': ref,
                    'p_authors': '',
                    'p_title': '',
                    'p_doi': ''
                }
                x1 = re.findall("(.*et al.?)([^.]+).", ref)
                x2 = re.findall("(.*)\d{4}.?\s?(\".*\")", ref)
                x3 = re.findall(r'doi:?\s?(.+).?', ref)

                if x3:
                    search_item['p_doi'] = x3[0].strip(punctuation).strip()

                if x1:
                    search_item['p_authors'] = x1[0][0].strip(punctuation).strip()
                    search_item['p_title'] = x1[0][1].strip(punctuation).strip()

                elif x2:
                    search_item['p_authors'] = x2[0][0].strip(punctuation).strip()
                    search_item['p_title'] = x2[0][1].strip(punctuation).strip()

                elif (ref.lower().startswith('who') or ref.lower().startswith('lci') \
                      or ref.lower().startswith('nvn') or ref.lower().startswith('nice')):
                    continue


            references_list.append(search_item)
    return references_list

def collectFromTables(rba_path):
    """Extracts the reference list from a file with a table

    Parameters
    ----------
    rba_path : string
        Path of the file to be processed

    Returns
    -------
    references_list : list of dictionaries
        Contains the extracted authors and title for each reference as an object
    """
    doc = Document(rba_path)

    data = []
    references_list = []
    for table in doc.tables:

        for i, row in enumerate(table.rows):

            for j, cell in enumerate(row.cells):

                if (cell.paragraphs[0].runs[0].bold
                    and ('samenvatting' in cell.text.lower()
                         or 'summary' in cell.text.lower())):

                    bold_text = ''
                    p_title = ''
                    p_authors = ''

                    for p, paragraph in enumerate(cell.paragraphs):

                        if len(paragraph.runs) < 1:
                            continue

                        if (p == 0
                            and paragraph.runs[0].bold
                            and not paragraph.runs[-1].bold):

                            for r, run in enumerate(paragraph.runs):

                                if run.bold:
                                    bold_text += run.text.strip() + ' '

                                else:
                                    p_title = bold_text

                                if p_title and not run.bold and run.text != ' ':

                                    p_authors += run.text.strip() + ' '

                        elif (paragraph.runs[0].bold
                              and paragraph.runs[-1].bold and p_title == ''):

                            bold_text += paragraph.text.strip() + ' '

                        elif (not paragraph.runs[-1].bold
                              and p > 0
                              and len(cell.paragraphs[p-1].runs) > 0
                              and cell.paragraphs[p-1].runs[-1].bold):

                            p_title = bold_text

                            p_authors += paragraph.text.strip() + ' '

                        elif (paragraph.runs[0].underline
                              and (paragraph.runs[0].text.lower().startswith('samenvatting')
                                   or paragraph.runs[0].text.lower().startswith('summary'))):
                            break

                    if p_title.strip() != ''  and re.search(r'^[0-9].[0-9]', p_title) is None:
                        references_list.append({
                            'p_title': p_title.replace('\xa0',' ').replace('\n', '').strip(),
                            'p_authors': p_authors.replace('\xa0',' ').replace('\n', '').strip(),
                        })

    return references_list

def pubmed2csv(references_list, csv_path):
    """Creates a csv file consists of the references

    Parameters
    ----------
    references_list : list of dictionaries
        Contains dictionaries with paper DOIs, authors and titles

    csv_path : string
        corresponding csv path for the file

    Returns
    -------
    Creates a csv file on the given CSV path
    """
    if not os.path.exists(csv_path):
        header = ['record_id', 'title', 'abstract', 'doi', 'final_included']
        with open(csv_path, 'w', newline='') as csvfile:
            cw = csv.writer(csvfile, delimiter=',')
            cw.writerow(header)

    #Loop over the reference list
    for ref in references_list:
        search_query = ref['p_doi'] if 'p_doi' in ref and len(ref['p_doi']) > 0 else ref['p_title']
        try:
            PubmedSearchResults = search_API(search_query)
            id_list = PubmedSearchResults['IdList']
            papers = fetch_details(id_list)
            referenceFound = False

            for paperIndex, paper in enumerate(papers['PubmedArticle']):
                paperTitle = paper['MedlineCitation']['Article']['ArticleTitle'].strip(punctuation).strip()
                if (Simhash(paperTitle).distance(Simhash(ref['p_title'])) <= 5):
                    doi = ''
                    for idn, ids in enumerate(papers['PubmedArticle'][paperIndex]['PubmedData']['ArticleIdList']):
                        if papers['PubmedArticle'][paperIndex]['PubmedData']['ArticleIdList'][idn].attributes['IdType'] == 'doi':
                            doi = papers['PubmedArticle'][paperIndex]['PubmedData']['ArticleIdList'][idn]

                    abstract = ''
                    if 'Abstract' in paper['MedlineCitation']['Article']:
                        abstract = ' '.join([str(x) for x in paper['MedlineCitation']['Article']['Abstract']['AbstractText']])

                    data = [id_list[paperIndex], paperTitle, abstract, doi, 1]
                    with open(csv_path, 'a') as csvfile:
                        cw = csv.writer(csvfile, delimiter=',')
                        cw.writerow(data)
                    referenceFound = True
                    break
            if not referenceFound:
                print('!!!!!!!!!!!!!!! The reference could not be found automaticaly: ', search_query)
        except:
            print('!!!! The reference could not be found automaticaly: ', ref['p_title'])
